from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

from .models import LocalizedVariant, ScreenplayEpisode, StoryDocument, StoryPackage


def write_story_exports(document: StoryDocument, output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = (
        f"{_safe_stem(document.package.title if document.package else document.dna.title)}"
        f"_{document.version}"
    )
    paths: list[Path] = []

    canonical_json = output_dir / f"{stem}_story_package.json"
    canonical_json.write_text(
        json.dumps(document.model_dump(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    paths.append(canonical_json)

    markdown = output_dir / f"{stem}_story_package.md"
    markdown.write_text(_document_markdown(document), encoding="utf-8")
    paths.append(markdown)

    if document.package is not None:
        fountain = output_dir / f"{stem}.fountain"
        fountain.write_text(_package_fountain(document.package), encoding="utf-8")
        paths.append(fountain)
        paths.extend(_screenplain_exports(fountain, output_dir))

        docx = output_dir / f"{stem}_story_package.docx"
        if _write_docx(document, docx):
            paths.append(docx)

        pdf = output_dir / f"{stem}_story_package.pdf"
        if _write_pdf(document, pdf):
            paths.append(pdf)

    if document.similarity_report is not None:
        report_path = output_dir / f"{stem}_similarity_report.json"
        report_path.write_text(
            json.dumps(
                document.similarity_report.model_dump(), ensure_ascii=False, indent=2
            ),
            encoding="utf-8",
        )
        paths.append(report_path)

    for variant in document.localizations:
        localized = (
            output_dir / f"{stem}_{variant.language}_{_safe_stem(variant.market)}.md"
        )
        localized.write_text(_localized_markdown(variant), encoding="utf-8")
        paths.append(localized)
    return paths


def _document_markdown(document: StoryDocument) -> str:
    package = document.package
    lines = [
        f"# {(package.title if package else document.dna.title) or 'Crayotter Story Package'}",
        "",
        f"- Version: `{document.version}`",
        f"- Status: `{document.status}`",
        f"- Content type: `{document.request.content_type}`",
        f"- Target markets: {', '.join(document.request.target_markets)}",
        f"- Source hashes: {', '.join(document.provenance.source_hashes.values())}",
        "",
        "## Story DNA",
        "",
        f"**Logline:** {document.dna.logline}",
        "",
        f"**Premise:** {document.dna.premise}",
        "",
        f"**Conflict engine:** {document.dna.conflict_engine}",
        "",
        f"**Emotional promise:** {document.dna.emotional_promise}",
        "",
        "### Characters",
        "",
    ]
    for character in document.dna.characters:
        lines.extend(
            [
                f"#### {character.name} (`{character.character_id}`)",
                "",
                f"- Role: {character.role}",
                f"- Goal: {character.goal}",
                f"- Motivation: {character.motivation}",
                f"- Arc: {character.arc}",
                "",
            ]
        )
    lines.extend(["### Episode structure", ""])
    for episode in document.dna.episodes:
        lines.extend(
            [
                f"#### Episode {episode.number}: {episode.title}",
                "",
                f"- Incoming: {episode.incoming_state}",
                f"- Objective: {episode.objective}",
                f"- Escalation: {episode.escalation}",
                f"- Turn: {episode.turn}",
                f"- Payoff: {episode.payoff}",
                f"- Outgoing pressure: {episode.outgoing_pressure}",
                "",
            ]
        )
    lines.extend(["## Alternative directions", ""])
    for direction in document.directions:
        marker = (
            " — **selected**"
            if direction.direction_id == document.selected_direction_id
            else ""
        )
        lines.extend(
            [
                f"### {direction.title}{marker}",
                "",
                direction.logline,
                "",
                f"- Strategy: {direction.protagonist_strategy}",
                f"- Opposition: {direction.opposition_mechanism}",
                f"- Cost: {direction.cost}",
                f"- Differentiation: {direction.differentiation}",
                "",
            ]
        )
    if package is not None:
        lines.extend(_package_markdown_lines(package))
    if document.localizations:
        lines.extend(["## Localized variants", ""])
        for variant in document.localizations:
            lines.extend(
                [
                    f"### {variant.language.upper()} · {variant.market}",
                    "",
                    variant.logline,
                    "",
                    *[f"- {item}" for item in variant.cultural_changes],
                    "",
                ]
            )
    if document.similarity_report is not None:
        report = document.similarity_report
        lines.extend(
            [
                "## Similarity risk",
                "",
                f"**Overall:** {report.overall_risk.upper()} ({report.overall_score:.3f})",
                "",
                report.disclaimer,
                "",
            ]
        )
        for finding in report.findings:
            lines.extend(
                [
                    f"### {finding.signal} · {finding.risk.upper()} · {finding.score:.3f}",
                    "",
                    f"- Output: `{finding.output_locator}` — {finding.output_excerpt}",
                    f"- Reference: `{finding.source_id}:{finding.source_locator}` — {finding.source_excerpt}",
                    f"- Recommendation: {finding.recommendation}",
                    "",
                ]
            )
    return "\n".join(lines).strip() + "\n"


def render_story_markdown(document: StoryDocument) -> str:
    """Public, side-effect-free Markdown renderer for approved-script handoffs."""

    return _document_markdown(document)


def render_production_screenplay(document: StoryDocument) -> str:
    """Render only approved production instructions, excluding ideation alternatives."""

    package = document.package
    if package is None:
        raise ValueError("A production screenplay requires a story package.")
    lines = [
        f"# {package.title or 'Crayotter Production Screenplay'}",
        "",
        f"- Approved version: `{document.version}`",
        "- Status: `APPROVED — LOCKED FOR VIDEO PRODUCTION`",
        f"- Content type: `{document.request.content_type}`",
        "",
        "## Locked production brief",
        "",
        f"**Logline:** {package.logline}",
        "",
        f"**Synopsis:** {package.synopsis}",
        "",
        f"**World / visual constraints:** {package.world}",
        "",
    ]
    for episode in package.episodes:
        lines.extend(
            [
                f"## Episode {episode.number} production target",
                "",
                f"- Target duration: {episode.target_duration_seconds} seconds",
                f"- Scene count: {len(episode.scenes)}",
                "",
            ]
        )
        lines.extend(_episode_markdown_lines(episode))
    lines.extend(["## Locked video prompts", ""])
    for item in package.video_prompt_package:
        lines.append(f"- `{item.get('scene_id', '')}`: {item.get('prompt', '')}")
    lines.extend(
        [
            "",
            "## Production rule",
            "",
            "Preserve the selected episode's scene order, dialogue, narration, ending, timing intent, "
            "and voice requirements. Material selection may adapt visuals, but must not rewrite the script.",
            "",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _package_markdown_lines(package: StoryPackage) -> list[str]:
    lines = [
        "## New story package",
        "",
        f"**Logline:** {package.logline}",
        "",
        f"**Synopsis:** {package.synopsis}",
        "",
        f"**World:** {package.world}",
        "",
    ]
    for episode in package.episodes:
        lines.extend(_episode_markdown_lines(episode))
    lines.extend(["## Video prompt package", ""])
    for item in package.video_prompt_package:
        lines.append(f"- `{item.get('scene_id', '')}`: {item.get('prompt', '')}")
    lines.append("")
    return lines


def _episode_markdown_lines(episode: ScreenplayEpisode) -> list[str]:
    lines = [f"### Episode {episode.number}: {episode.title}", "", episode.synopsis, ""]
    for scene in episode.scenes:
        lines.extend(
            [
                f"#### {scene.number}. {scene.heading}",
                "",
                scene.action,
                "",
            ]
        )
        for dialogue in scene.dialogue:
            character = dialogue.get("character", "CHARACTER")
            direction = dialogue.get("direction", "")
            lines.extend(
                [
                    f"**{character}**{f' ({direction})' if direction else ''}",
                    "",
                    dialogue.get("text", ""),
                    "",
                ]
            )
        if scene.hook:
            lines.extend([f"*Hook: {scene.hook}*", ""])
    return lines


def _package_fountain(package: StoryPackage) -> str:
    lines = [
        f"Title: {package.title}",
        "Credit: Generated with Crayotter",
        "Draft date: Automated development draft",
        "",
    ]
    for episode in package.episodes:
        lines.extend(
            [
                f"# EPISODE {episode.number}: {episode.title}",
                f"= {episode.synopsis}",
                "",
            ]
        )
        for scene in episode.scenes:
            heading = (
                scene.heading
                or f"INT. {scene.location or 'LOCATION'} - {scene.time_of_day or 'DAY'}"
            )
            lines.extend([heading.upper(), "", scene.action, ""])
            for dialogue in scene.dialogue:
                character = dialogue.get("character", "CHARACTER").upper()
                lines.append(character)
                if dialogue.get("direction"):
                    lines.append(f"({dialogue['direction']})")
                lines.extend([dialogue.get("text", ""), ""])
            if scene.hook:
                lines.extend([f"[[HOOK: {scene.hook}]]", ""])
    return "\n".join(lines).strip() + "\n"


def _screenplain_exports(fountain: Path, output_dir: Path) -> list[Path]:
    executable = shutil.which("screenplain")
    if not executable:
        return []
    created: list[Path] = []
    for suffix in (".html", ".fdx"):
        target = output_dir / f"{fountain.stem}{suffix}"
        try:
            completed = subprocess.run(
                [executable, str(fountain), str(target)],
                check=False,
                capture_output=True,
                text=True,
                timeout=60,
            )
            if completed.returncode == 0 and target.is_file():
                created.append(target)
        except (OSError, subprocess.SubprocessError):
            continue
    return created


def _write_docx(document: StoryDocument, destination: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        return False
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10.5)
        package = document.package
        doc.add_heading(package.title if package else document.dna.title, level=0)
        doc.add_paragraph(f"Version {document.version} · {document.status}")
        doc.add_heading("Story DNA", level=1)
        doc.add_paragraph(document.dna.logline)
        doc.add_paragraph(document.dna.conflict_engine)
        doc.add_heading("Characters", level=2)
        for character in document.dna.characters:
            doc.add_heading(character.name, level=3)
            doc.add_paragraph(f"{character.role} · {character.goal} · {character.arc}")
        if package is not None:
            doc.add_heading("New story package", level=1)
            doc.add_paragraph(package.synopsis)
            for episode in package.episodes:
                doc.add_heading(f"Episode {episode.number}: {episode.title}", level=2)
                doc.add_paragraph(episode.synopsis)
                for scene in episode.scenes:
                    doc.add_heading(f"{scene.number}. {scene.heading}", level=3)
                    doc.add_paragraph(scene.action)
                    for dialogue in scene.dialogue:
                        paragraph = doc.add_paragraph()
                        paragraph.add_run(
                            dialogue.get("character", "CHARACTER")
                        ).bold = True
                        paragraph.add_run(f"\n{dialogue.get('text', '')}")
        if document.similarity_report is not None:
            doc.add_heading("Similarity risk", level=1)
            report = document.similarity_report
            doc.add_paragraph(
                f"{report.overall_risk.upper()} · {report.overall_score:.3f}"
            )
            doc.add_paragraph(report.disclaimer)
            for finding in report.findings:
                doc.add_heading(f"{finding.signal} · {finding.risk}", level=3)
                doc.add_paragraph(finding.explanation)
                doc.add_paragraph(f"Recommendation: {finding.recommendation}")
        doc.save(destination)
        return destination.is_file()
    except Exception:
        return False


def _write_pdf(document: StoryDocument, destination: Path) -> bool:
    try:
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        return False
    try:
        root = Path(__file__).resolve().parents[2]
        font_path = root / "assets" / "fonts" / "AlibabaPuHuiTi-3-55-Regular.ttf"
        font_name = "CrayotterCJK"
        if font_path.is_file():
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
        else:
            font_name = "Helvetica"
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "CrayotterBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=9.5,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=5,
            wordWrap="CJK",
        )
        heading = ParagraphStyle(
            "CrayotterHeading",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=15,
            leading=19,
            spaceBefore=10,
            spaceAfter=6,
            wordWrap="CJK",
        )
        title = ParagraphStyle(
            "CrayotterTitle",
            parent=heading,
            fontSize=22,
            leading=27,
            spaceAfter=14,
        )
        package = document.package
        story = [
            Paragraph(_xml(package.title if package else document.dna.title), title)
        ]
        story.extend(
            [
                Paragraph(
                    _xml(f"Version {document.version} · {document.status}"), body
                ),
                Paragraph("Story DNA", heading),
                Paragraph(_xml(document.dna.logline), body),
                Paragraph(_xml(document.dna.conflict_engine), body),
                Spacer(1, 4 * mm),
            ]
        )
        if package is not None:
            story.extend(
                [
                    Paragraph("New story package", heading),
                    Paragraph(_xml(package.synopsis), body),
                ]
            )
            for episode in package.episodes:
                story.append(
                    Paragraph(
                        _xml(f"Episode {episode.number}: {episode.title}"), heading
                    )
                )
                story.append(Paragraph(_xml(episode.synopsis), body))
                for scene in episode.scenes:
                    story.append(
                        Paragraph(_xml(f"{scene.number}. {scene.heading}"), heading)
                    )
                    story.append(Paragraph(_xml(scene.action), body))
                    for dialogue in scene.dialogue:
                        text = f"<b>{_xml(dialogue.get('character', 'CHARACTER'))}</b><br/>{_xml(dialogue.get('text', ''))}"
                        story.append(Paragraph(text, body))
        if document.similarity_report is not None:
            story.extend(
                [
                    PageBreak(),
                    Paragraph("Similarity risk", heading),
                    Paragraph(
                        _xml(
                            f"{document.similarity_report.overall_risk.upper()} · "
                            f"{document.similarity_report.overall_score:.3f}"
                        ),
                        body,
                    ),
                    Paragraph(_xml(document.similarity_report.disclaimer), body),
                ]
            )
        SimpleDocTemplate(
            str(destination),
            pagesize=A4,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=package.title if package else document.dna.title,
        ).build(story)
        return destination.is_file()
    except Exception:
        return False


def _localized_markdown(variant: LocalizedVariant) -> str:
    lines = [
        f"# {variant.title}",
        "",
        f"Language: `{variant.language}` · Market: `{variant.market}`",
        "",
        variant.logline,
        "",
        variant.synopsis,
        "",
        "## Cultural changes",
        "",
        *[f"- {item}" for item in variant.cultural_changes],
        "",
    ]
    for episode in variant.episodes:
        lines.extend(_episode_markdown_lines(episode))
    return "\n".join(lines).strip() + "\n"


def _safe_stem(value: str) -> str:
    cleaned = "".join(
        char if char.isalnum() or char in {"-", "_"} else "_" for char in value
    )
    return cleaned.strip("_")[:80] or "crayotter_story"


def _xml(value: str) -> str:
    return (
        str(value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )
